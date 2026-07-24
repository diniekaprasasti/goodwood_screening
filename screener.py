# -*- coding: utf-8 -*-
"""
Goodwood Screening Engine
Mengecek kelengkapan struktur manuskrip (.docx) sebelum masuk proses review.
"""
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

# ---------------------------------------------------------------
# Konfigurasi kriteria (mudah diubah jika kebijakan jurnal berubah)
# ---------------------------------------------------------------
TITLE_MAX_WORDS = 15
KEYWORDS_MIN, KEYWORDS_MAX = 3, 5
ABSTRACT_MIN, ABSTRACT_MAX = 180, 250
REF_MIN = 30
REF_INTL_PCT = 80.0
DOC_WORDS_MIN, DOC_WORDS_MAX = 5000, 10000

CHAPTERS = [
    ("Introduction", [r"introduction", r"pendahuluan"]),
    ("Literature Review & Hypothesis Development",
     [r"literature\s+review", r"hypothes[ei]s\s+development", r"tinjauan\s+pustaka"]),
    ("Research Methodology",
     [r"research\s+method(?:ology|s)?", r"methodology", r"\bmethods?\b", r"metod[eo]"]),
    ("Result and Discussion",
     [r"results?\s*(?:and|&|dan)\s*discussions?", r"findings?\s*(?:and|&)\s*discussions?",
      r"hasil\s+dan\s+pembahasan", r"results?\s*,?\s*(?:analysis|and\s+findings?)",
      r"^results?$", r"^discussions?$", r"^findings?$", r"^hasil$", r"^pembahasan$",
      r"^empirical\s+results?$", r"^analysis\s+and\s+discussions?$"]),
    ("Conclusion", [r"conclusions?", r"concluding\s+remarks", r"kesimpulan"]),
]

# Template jurnal Goodwood mewajibkan Bab 5 (Conclusion) dipecah menjadi TIGA
# subsection berikut. Subsection "Conclusion" dianggap otomatis terpenuhi oleh
# heading Bab 5 itu sendiri (mis. "5. Conclusion" atau "5.1 Conclusion") —
# yang divalidasi keberadaannya secara eksplisit hanya dua subsection lainnya.
CONCLUSION_SUBSECTIONS = [
    ("Conclusion", [r"^conclusions?$", r"^concluding\s+remarks$", r"^kesimpulan$"]),
    ("Research Limitations",
     [r"research\s+limitations?", r"limitations?\s+of\s+(?:the\s+)?(?:study|research)",
      r"study\s+limitations?", r"keterbatasan\s+penelitian"]),
    ("Suggestions and Directions for Future Research",
     [r"suggestions?\s+and\s+directions?\s+for\s+future\s+research",
      r"directions?\s+for\s+future\s+research", r"future\s+research\s+(?:directions?|agenda)",
      r"recommendations?\s+for\s+future\s+research",
      r"saran\s+(?:dan\s+arah\s+)?penelitian\s+(?:selanjutnya|mendatang|masa\s+depan)"]),
]

SPECIAL_CHAPTERS = [
    ("Acknowledgement", [r"acknowledg?e?ments?"]),
    ("Author Contribution",
     [r"authors?['\u2019]?s?\s*contributions?", r"contributions?\s+of\s+(?:the\s+)?authors?",
      r"credit\s+authorship", r"authorship\s+contributions?", r"kontribusi\s+penulis"]),
]

REFERENCE_HEADINGS = [r"^references?\b", r"^bibliography\b", r"^daftar\s+pustaka\b"]

# Elemen ilmiah non-heading: caption Table/Figure/Formula/Chart/Appendix, dsb.
# Paragraf yang cocok pola ini TIDAK PERNAH dianggap sebagai heading struktur
# dokumen (Section/Subsection/Sub-subsection), meskipun teksnya pendek & bold.
CAPTION_PATTERNS = [
    r"^table\s+\d", r"^tabel\s+\d",
    r"^figure\s+\d", r"^fig\.?\s+\d", r"^gambar\s+\d",
    r"^equation\s+\d", r"^eq\.?\s+\d", r"^persamaan\s+\d", r"^rumus\s+\d",
    r"^chart\s+\d", r"^grafik\s+\d", r"^diagram\s+\d",
    r"^appendix\s+\d", r"^lampiran\s+\d",
]

# Pola caption Table/Figure KHUSUS untuk validasi cross-reference — memisahkan
# jenis elemen (Table/Figure) beserta NOMORNYA, supaya bisa dicocokkan dengan
# kemunculan rujukan yang sama di badan teks (mis. "Table 2" harus disebut
# lagi di luar caption-nya sendiri).
CROSSREF_ELEMENTS = [
    ("Table", ["table", "tabel"], [r"^table\s+(\d+)", r"^tabel\s+(\d+)"]),
    ("Figure", ["figure", "fig", "gambar"], [r"^figure\s+(\d+)", r"^fig\.?\s+(\d+)", r"^gambar\s+(\d+)"]),
]

# Pernyataan hipotesis (mis. "H1: X1 → Y", "H2. ...", "Hypothesis 3: ...") —
# bukan heading struktur dokumen, meskipun sering ditulis pendek/bold/italic
# dan (karena hanya berisi huruf notasi seperti H/X/Y) bisa lolos heuristik
# "semua huruf kapital".
HYPOTHESIS_PATTERNS = [
    r"^h\d+\s*[:.\)]", r"^ha\d+\s*[:.\)]", r"^h0\s*[:.\)]",
    r"^hypothes[ei]s\s*\d+\s*[:.\)]", r"^hipotesis\s*\d+\s*[:.\)]",
]

# Kata penanda referensi berbahasa/terbitan Indonesia (jurnal nasional / non-internasional)
ID_MARKERS = [
    r"\bjurnal\b", r"\buniversitas\b", r"\bskripsi\b", r"\btesis\b", r"\bdisertasi\b",
    r"\bpenerbit\b", r"\bfakultas\b", r"\bekonomi\s+dan\b", r"\bakuntansi\s+dan\b",
    r"\bmanajemen\s+dan\b", r"\bindonesia\b", r"\bjakarta\b", r"\byogyakarta\b",
    r"\bbandung\b", r"\bsemarang\b", r"\bsurabaya\b", r"\bterhadap\b", r"\bpengaruh\b",
    r"\banalisis\b", r"\bdan\b",
]

# Penanda referensi berupa BUKU / prosiding / web (bukan artikel jurnal)
BOOK_MARKERS = [
    r"\bpress\b", r"\bpublishing\b", r"\bpublishers?\b", r"\bedition\b", r"\bed\.\b",
    r"\bwiley\b", r"\bspringer\b(?!.*journal)", r"\broutledge\b", r"\bmcgraw", r"\bpearson\b",
    r"\bsage publications\b", r"\bhandbook\b", r"\btextbook\b", r"\bnew york:\s", r"\blondon:\s",
    r"\bboston:\s", r"\bcengage\b", r"\bprentice\s+hall\b",
]
PROC_MARKERS = [r"\bproceedings?\b", r"\bconference\b", r"\bsymposium\b", r"\bseminar\b"]
WEB_MARKERS = [r"\bretrieved\s+from\b", r"\baccessed\b", r"\bavailable\s+at\b", r"\bdiakses\b"]

# Penanda preprint (belum melalui peer-review formal di jurnal/prosiding) —
# mis. arXiv, SSRN, bioRxiv, dsb. Referensi jenis ini TIDAK memiliki struktur
# "Nama Jurnal, Vol(Issue), Halaman" sehingga sebelumnya sering jatuh ke
# kategori "Lainnya/Tidak terdeteksi" dan luput dari validasi.
PREPRINT_MARKERS = [
    r"\barxiv\b", r"\bssrn\b", r"\bbiorxiv\b", r"\bmedrxiv\b", r"\bresearchgate\b",
    r"\bpreprints?\.org\b", r"\bpreprint\b",
]

# Penanda kuat artikel jurnal
JOURNAL_CUES = [
    r"\bjournal\b", r"\breview of\b", r"\bquarterly\b", r"\bdoi\b", r"doi\.org",
    r"\d+\s*\(\d+[\-–]?\d*\)\s*[,:]?\s*(?:pp\.?\s*)?\d+\s*[\-–]\s*\d+",  # 12(3), 45-67
    r"\bvol\.?\s*\d+", r"\bno\.?\s*\d+", r"\bpp\.?\s*\d+\s*[\-–]\s*\d+",
]

# Pola DOI: "10.xxxx/xxxxx" ataupun URL "https://doi.org/10.xxxx/xxxxx"
DOI_PATTERN = re.compile(r"(?:doi\s*[:.]?\s*)?(10\.\d{4,9}/\S+)", re.I)

# ---------------------------------------------------------------
# Daftar jurnal terbitan Goodwood Publishing — dipakai untuk memvalidasi
# apakah manuskrip sudah menyitasi minimal N artikel dari jurnal-jurnal ini
# (lihat GOODWOOD_JOURNAL_MIN & validate_goodwood_journal_usage).
# ---------------------------------------------------------------
GOODWOOD_JOURNALS_RAW = [
    "Annals of Human Resource Management Research (AHRMR)",
    "Annals of Management and Organization Research (AMOR)",
    "Reviu Akuntansi, Manajemen, dan Bisnis (RAMBIS)",
    "Jurnal Akuntansi, Keuangan, dan Manajemen (Jakman)",
    "Studi Ilmu Manajemen dan Organisasi (SIMO)",
    "Studi Akuntansi, Keuangan, dan Manajemen (Sakman)",
    "Jurnal Ilmiah Hukum dan Hak Asasi Manusia (JIHHAM)",
    "Yumary: Jurnal Pengabdian kepada Masyarakat",
    "Relevansi: Stie Krakatau",
    "International Journal of Financial, Accounting, and Management (IJFAM)",
    "Journal of Social, Humanity, and Education (JSHE)",
    "Journal of Governance and Accountability Studies (JGAS)",
    "Journal of Sustainable Tourism and Entrepreneurship (JoSTE)",
    "Goodwood Akuntansi dan Auditing Reviu",
    "Jurnal Studi Perhotelan dan Pariwisata (JSPP)",
    "Jurnal Ilmu Siber dan Teknologi Digital (JISTED)",
    "Jurnal Studi Ilmu Sosial dan Politik (Jasispol)",
    "Kajian Psikologi dan Kesehatan Mental (KPKM)",
    "Jurnal Bisnis dan Pemasaran Digital (JBPD)",
    "Jurnal Nusantara Mengabdi (JNM)",
    "Jurnal Studi Pemerintahan dan Akuntabilitas (Jastaka)",
    "Kajian Ilmiah Hukum dan Kenegaraan (KIHAN)",
    "Jurnal Humaniora dan Ilmu Pendidikan (Jahidik)",
    "Bukhori: Kajian Ekonomi dan Keuangan Islam",
    "Jurnal Pemberdayaan Umat",
    "Jurnal Pemberdayaan Ekonomi",
    "Ners Akademika",
    "Jurnal Abdimas Multidisiplin (JAMU)",
    "Jurnal Ilmu Medis Indonesia (JIMI)",
    "Studi Ekonomi dan Kebijakan Publik (SEKP)",
    "Jurnal Kesehatan dan Keselamatan Kerja",
    "Jurnal Studi Multidisiplin Ilmu (Jasmi)",
    "Jurnal Teknologi Riset Terapan (JATRA)",
    "Jurnal Sistem Informasi Akuntansi dan Manajemen (Jusiam)",
    "Jurnal Kesehatan Maternal dan Neonatal (JESMAN)",
    "Jurnal Ilmiah Pertanian dan Peternakan (Jipper)",
    "Review of Nursing and Healthcare Research (RNHR)",
    "Jurnal Internasional Sistem Informasi Akuntansi dan Manajemen (IJAMIS)",
    "Journal of Multidisciplinary Academic and Practice Studies",
    "Journal of Multidisciplinary Academic Business Studies",
    "Psikohealth: Jurnal Ilmiah Psikologi dan Kesehatan Mental",
    "Dynamics of Politics and Democracy",
    "Journal of Digital Business and Marketing (JDBM)",
    "Annals of Animal Studies (AAS)",
    "Dirham: Journal of Sharia Finance and Economics (JoSFE)",
    "Annals of Justice and Humanity (AJH)",
    "Annals of Sustainable Agriculture and Forestry (ASAF)",
    "Jurnal Ilmu Metabolik dan Olahraga",
    "Studies in Medicine and Public Health (SiMPH)",
    "Studies in Economy and Public Policy (SEPP)",
    "Jurnal Linguistik dan Sastra Universal (ULLJ)",
    "Jurnal Ilmiah Manajemen Halal",
    "Jurnal Ilmu Fisiologi dan Biomedis",
    "Journal of Indigenous Culture, Tourism, and Language (JICTL)",
    "Al-Qadha: Jurnal Hukum Islam (AJHI)",
    "Jurnal Kajian Hukum Kontemporer (JKHK)",
    "Jurnal Kecerdasan Buatan dan Pembelajaran Mesin (JKBPM)",
    "Applied AI and Machine Learning Journal (AIML Journal)",
    "Advances in Financial Crime and Law (AFCL)",
    "Advances in Management and Business Studies (AMBuS)",
    "Advances in Public Law and Policy (AiPLAP)",
    "Global Academy of Business Studies (GABS)",
    "Global Academy of Multidisciplinary Studies (GAMS)",
    "International Journal of Higher Education Policy and Management (IJHEPM)",
    "Journal of Gender and Organizational Leadership (JGOL)",
]

GOODWOOD_JOURNAL_MIN = 5


def _normalize_journal_name(name):
    """Normalisasi nama jurnal untuk pencocokan: lowercase, buang tanda baca,
    rapikan spasi. Dipakai baik untuk daftar Goodwood maupun nama jurnal hasil
    ekstraksi referensi, supaya variasi tanda baca/kapitalisasi kecil (mis.
    "Reviu Akuntansi, Manajemen, dan Bisnis" vs "Reviu Akuntansi Manajemen dan
    Bisnis") tetap dianggap cocok."""
    name = name.lower()
    name = re.sub(r"[^a-z0-9\s]", " ", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name


def _build_goodwood_index():
    full_list, abbr_list = [], []
    for entry in GOODWOOD_JOURNALS_RAW:
        m = re.match(r"^(.*?)\s*\(([^)]+)\)\s*$", entry)
        if m:
            full, abbr = m.group(1).strip(), m.group(2).strip()
        else:
            full, abbr = entry.strip(), None
        full_norm = _normalize_journal_name(full)
        if full_norm:
            full_list.append((full_norm, entry))
        if abbr:
            abbr_norm = _normalize_journal_name(abbr)
            if abbr_norm:
                abbr_list.append((abbr_norm, entry))
    return full_list, abbr_list


_GOODWOOD_FULL_INDEX, _GOODWOOD_ABBR_INDEX = _build_goodwood_index()


def match_goodwood_journal(journal_name):
    """
    Kembalikan nama resmi jurnal Goodwood Publishing (lengkap dengan
    singkatannya, sesuai daftar GOODWOOD_JOURNALS_RAW) jika ``journal_name``
    cocok dengan salah satu jurnal Goodwood, atau None jika tidak cocok.

    Pencocokan bertahap dari yang paling ketat ke paling longgar: (1) sama
    persis (dengan normalisasi kapitalisasi/tanda baca), (2) cocok dengan
    singkatan resmi jurnal (mis. referensi hanya menulis "SIMO"), (3) salah
    satu nama memuat nama yang lain (toleransi subjudul/variasi kecil), lalu
    (4) kemiripan string (fuzzy match) untuk menoleransi typo/variasi
    penulisan kecil lainnya.
    """
    if not journal_name:
        return None
    norm = _normalize_journal_name(journal_name)
    if not norm or len(norm) < 3:
        return None

    for full_norm, original in _GOODWOOD_FULL_INDEX:
        if norm == full_norm:
            return original
    for abbr_norm, original in _GOODWOOD_ABBR_INDEX:
        if norm == abbr_norm:
            return original
    for full_norm, original in _GOODWOOD_FULL_INDEX:
        if len(full_norm) >= 8 and (full_norm in norm or norm in full_norm):
            return original

    import difflib
    best_ratio, best_match = 0.0, None
    for full_norm, original in _GOODWOOD_FULL_INDEX:
        ratio = difflib.SequenceMatcher(None, norm, full_norm).ratio()
        if ratio > best_ratio:
            best_ratio, best_match = ratio, original
    if best_ratio >= 0.87:
        return best_match
    return None


def validate_goodwood_journal_usage(ref_rows):
    """
    Memeriksa berapa banyak referensi pada daftar pustaka yang berasal dari
    jurnal terbitan Goodwood Publishing (lihat GOODWOOD_JOURNALS_RAW).
    Mengembalikan list of dict {"no", "text", "matched_journal"} — satu
    entri per referensi yang cocok.
    """
    matches = []
    for row in ref_rows:
        jn = row.get("journal_name")
        if not jn:
            continue
        official = match_goodwood_journal(jn)
        if official:
            matches.append({"no": row["no"], "text": row["text"], "matched_journal": official})
    return matches

# ---------------------------------------------------------------
# Pola-pola tambahan untuk validasi rinci format APA 6th Edition per-referensi
# ---------------------------------------------------------------
YEAR_RE = re.compile(r"\((\d{4}[a-z]?)(?:,\s*[^()]{0,40})?\)")

# Referensi web/blog/berita tanpa tahun terbit yang jelas sering ditulis
# dengan penanda "(n.d.)" ("no date") sesuai APA 6th Edition, mis.
# "Google Maps. (n.d.). Retrieved ... from ...". Penanda ini BUKAN
# kesalahan format — jadi perlu dikenali secara terpisah dari YEAR_RE agar
# tidak salah di-flag sebagai "tahun publikasi tidak ditemukan", dan agar
# sitasi in-text yang memakai tahun akses (mis. tahun retrieval) terhadap
# referensi semacam ini tetap bisa dicocokkan berdasarkan nama penulis saja.
NO_DATE_RE = re.compile(r"\(\s*n\.?\s*d\.?\s*\)", re.I)

# Nama jurnal + volume(issue) + halaman/nomor artikel, mis. ". Journal of X,
# 12(3), 45-67". Disusun sebagai BEBERAPA pola dari yang paling spesifik ke
# paling longgar, supaya referensi yang TIDAK memiliki rentang halaman (mis.
# hanya nomor artikel tunggal, atau tidak mencantumkan halaman sama sekali)
# tetap bisa diketahui nama jurnalnya — kekurangan volume/issue/halaman itu
# sendiri sudah dicek & dilaporkan terpisah oleh VOL_ISSUE_RE/PAGES_RE, jadi
# tidak perlu "menyandera" ekstraksi nama jurnal pada kelengkapan elemen lain.
JOURNAL_SEGMENT_PATTERNS = [
    # 1) Jurnal, Vol(Issue), halaman-rentang — mis. "Journal of X, 12(3), 45-67"
    re.compile(r"\.\s*([A-Z][^.]*?),\s*\d+\s*(?:\(\d+[a-zA-Z]?\))?\s*,?\s*\d+\s*[\-\u2013]\s*\d+"),
    # 2) Jurnal, Vol(Issue), nomor artikel/halaman tunggal — mis.
    #    "BMC Health Services Research, 25(1), 985" atau
    #    "Journal of Nursing Management, 2025(1), 6893336"
    re.compile(r"\.\s*([A-Z][^.]*?),\s*\d+\s*\(\d+[a-zA-Z]?\)\s*,\s*[a-zA-Z]?\d+\b"),
    # 3) Jurnal, Vol(Issue) — tanpa halaman/nomor artikel sama sekali, mis.
    #    "Journal of Ecohumanism, 3(8)."
    re.compile(r"\.\s*([A-Z][^.]*?),\s*\d+\s*\(\d+[a-zA-Z]?\)\s*[.,]?"),
    # 4) Jurnal, Vol (tanpa issue), halaman-rentang — mis. "Journal of X, 12, 45-67"
    re.compile(r"\.\s*([A-Z][^.]*?),\s*\d+\s*,\s*\d+\s*[\-\u2013]\s*\d+"),
    # 5) Jurnal, Vol (tanpa issue), nomor artikel/halaman tunggal — mis.
    #    "Decision Support Systems, 176, 114051" (jurnal dgn continuous
    #    pagination/article number, tanpa nomor issue sama sekali)
    re.compile(r"\.\s*([A-Z][^.]*?),\s*\d+\s*,\s*[a-zA-Z]?\d+\b"),
]


def _find_journal_segment(text):
    """Coba tiap pola JOURNAL_SEGMENT_PATTERNS berurutan (paling spesifik dulu)
    dan kembalikan match pertama yang cocok, atau None jika semuanya gagal."""
    for pat in JOURNAL_SEGMENT_PATTERNS:
        m = pat.search(text)
        if m:
            return m
    return None


VOL_ISSUE_RE = re.compile(r"\b\d{1,4}\s*\(\s*\d+[a-zA-Z]?\s*\)")
PAGES_RE = re.compile(r"\b(?:pp\.?\s*)?\d{1,5}\s*[\-\u2013]\s*\d{1,5}\b")
URL_NOT_DOI_RE = re.compile(r"https?://(?!(?:dx\.)?doi\.org)\S+", re.I)
# Awal segmen penulis APA: "NamaBelakang, I." atau "NamaBelakang II, I. I."
AUTHOR_START_RE = re.compile(
    r"^[A-Z][A-Za-z\u00c0-\u00ff'\-]+(?:\s[A-Z][A-Za-z\u00c0-\u00ff'\-]+)*,\s*[A-Z]\."
)
# Heuristik penulis lembaga/organisasi (mis. "World Health Organization",
# "Google Maps.") — tanpa inisial bergaya "I.", tidak perlu di-flag sebagai
# salah format. Titik di akhir bersifat OPSIONAL dan diperbolehkan di sini
# (beda dengan AUTHOR_START_RE) karena untuk nama lembaga titik tersebut
# adalah tanda akhir kalimat biasa, bukan inisial — sehingga baik dengan
# maupun tanpa titik penutup tetap dianggap format yang sah.
ORG_AUTHOR_RE = re.compile(r"^[A-Z][A-Za-z]*(?:\s+[A-Z&][A-Za-z]*){1,6}\.?$")
# Heuristik penulis bernama tunggal / mononim (mis. "Sugiyono", umum pada
# penulis Indonesia) — satu kata diawali huruf kapital, boleh diikuti titik,
# dan boleh diikuti penulis mononim lain yang digabung dengan "," atau "&"
# (mis. "Sugiyono & Kuncoro."). Sengaja HANYA menerima kata tunggal per nama
# (tanpa spasi di dalamnya) supaya nama depan+belakang yang salah format
# tanpa koma (mis. "John Smith") tetap terdeteksi sebagai kesalahan format.
SINGLE_NAME_AUTHOR_RE = re.compile(
    r"^[A-Z][a-zA-Z\u00c0-\u00ff'\-]{1,30}\.?"
    r"(?:\s*(?:,\s*&|&|,)\s*[A-Z][a-zA-Z\u00c0-\u00ff'\-]{1,30}\.?)*$"
)


def _is_title_case_keyword(kw):
    """
    True jika SETIAP kata pada sebuah keyword diawali huruf kapital (format
    "Capital Each Word" / Title Case), mis. "Service Quality", "PLS-SEM",
    "E-Commerce". Kata disini dipecah per spasi/tanda hubung/garis miring.
    Token yang sama sekali tidak memiliki huruf (mis. angka atau simbol
    murni) dilewati karena tidak relevan untuk dicek kapitalisasinya.
    """
    words = re.split(r"[\s\-/]+", kw.strip())
    for w in words:
        m = re.search(r"[A-Za-z]", w)
        if not m:
            continue
        if not w[m.start()].isupper():
            return False
    return True


def _match_any(text, patterns):
    t = text.lower()
    return any(re.search(p, t) for p in patterns)


def _wc(text):
    return len(re.findall(r"\S+", text.strip()))


def has_doi(text):
    return bool(DOI_PATTERN.search(text))


def needs_doi(ref_type):
    """Referensi jenis artikel jurnal (nasional/internasional) semestinya punya DOI."""
    return ref_type in ("Jurnal Internasional", "Jurnal Nasional")


# ---------------------------------------------------------------
# Ekstraksi paragraf dokumen — termasuk paragraf di dalam tabel
# ---------------------------------------------------------------
def _iter_block_items(parent):
    """Yield tiap child (paragraf/tabel) dari *parent* sesuai urutan dokumen."""
    if hasattr(parent, "element"):
        parent_elm = parent.element.body
    elif hasattr(parent, "_tc"):
        parent_elm = parent._tc
    else:
        parent_elm = parent._element
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def get_all_paragraphs(doc):
    """
    Mengembalikan SEMUA paragraf dalam dokumen secara berurutan, termasuk
    paragraf yang berada di dalam sel tabel (dan tabel bersarang di dalamnya).

    PENTING: banyak template jurnal menempatkan blok judul/abstrak/keywords
    di dalam tabel (mis. kolom kiri berisi logo & article history, kolom
    kanan berisi judul + abstrak). ``Document.paragraphs`` bawaan python-docx
    HANYA membaca paragraf level tubuh dokumen dan TIDAK menelusuri isi
    tabel, sehingga Abstract/Keywords yang sebenarnya ada bisa gagal
    terdeteksi (false negative). Fungsi ini memperbaiki hal tersebut.
    """
    paras = []
    in_table_flags = []  # sejajar dgn `paras`: True jika paragraf ada di dalam sel tabel

    def walk(parent, in_table):
        for block in _iter_block_items(parent):
            if isinstance(block, Paragraph):
                paras.append(block)
                in_table_flags.append(in_table)
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        walk(cell, True)

    walk(doc, False)
    return paras, in_table_flags


def _has_auto_numbering(para):
    """
    True jika paragraf memakai automatic list numbering Word (numPr) —
    baik didefinisikan langsung pada paragraf maupun diwariskan dari gaya
    (style) yang dipakainya. Banyak template jurnal me-link gaya "Heading 2"
    / "Heading 3" ke Multilevel List, sehingga nomor seperti "2.5" muncul
    otomatis saat dibuka di Word TAPI TIDAK tersimpan sebagai teks pada
    ``paragraph.text`` (python-docx hanya membaca run text, bukan hasil
    render numbering). Tanpa pengecekan ini, heading yang sebenarnya SUDAH
    bernomor benar bisa keliru terdeteksi sebagai "belum ada penomoran".
    """
    def _find_numpr(el):
        if el is None:
            return None
        pPr = el.find(qn("w:pPr"))
        if pPr is None:
            return None
        return pPr.find(qn("w:numPr"))

    if _find_numpr(para._p) is not None:
        return True
    style = para.style
    if style is not None and getattr(style, "element", None) is not None:
        if _find_numpr(style.element) is not None:
            return True
    return False


def _is_caption(text):
    """True jika teks adalah caption Table/Figure/Formula/Chart/Appendix, dsb,
    ATAU pernyataan hipotesis (mis. "H1: X1 -> Y"). Elemen semacam ini bukan
    bagian dari struktur Section/Subsection/Sub-subsection dokumen dan tidak
    boleh divalidasi sebagai heading."""
    t = text.strip()
    return _match_any(t, CAPTION_PATTERNS) or _match_any(t, HYPOTHESIS_PATTERNS)


def _is_headingish(para, in_table=False):
    """Deteksi paragraf yang berperan sebagai judul bab (heading struktur
    dokumen), BUKAN elemen ilmiah lain seperti caption Table/Figure/Formula,
    pernyataan hipotesis (H1/H2/...), atau isi sel tabel data (mis. header
    kolom "No." / "Criterion")."""
    style = (para.style.name or "").lower() if para.style else ""
    text = para.text.strip()
    if not text:
        return False
    # Caption Table/Figure/Formula/Chart/Appendix maupun pernyataan hipotesis
    # — selalu diabaikan, meskipun kebetulan diberi gaya Heading atau
    # full-bold/full-italic di template penulis.
    if _is_caption(text):
        return False
    if "heading" in style or style == "title":
        return True
    # Paragraf di dalam sel tabel (data tabel, header kolom, dsb.) bukan
    # bagian dari struktur heading dokumen kecuali memang diberi gaya
    # Heading eksplisit (sudah ditangani di atas).
    if in_table:
        return False
    # Teks pendek + bold penuh sering dipakai sebagai heading manual
    if _wc(text) <= 10:
        runs = [r for r in para.runs if r.text.strip()]
        if runs and all(r.bold for r in runs):
            return True
        # Heuristik "semua huruf kapital" HANYA berlaku jika teksnya memang
        # didominasi huruf (bukan notasi/rumus seperti "H1: X1 -> Y" yang
        # kebetulan huruf-hurufnya berupa variabel kapital H/X/Y).
        letters = [ch for ch in text if ch.isalpha()]
        if text.isupper() and len(text) > 3 and len(letters) >= 4:
            return True
    return False


def _heading_style_level(para):
    """
    Level heading berdasarkan gaya paragraf Word:
      1 = Heading 1 / Title, 2 = Heading 2, 3 = Heading 3+, 0 = tanpa gaya heading
      (mis. bold manual) — level akan ditentukan lewat heuristik lain.
    """
    style = (para.style.name or "").lower() if para.style else ""
    m = re.search(r"heading\s*(\d+)", style)
    if m:
        return min(int(m.group(1)), 3)
    if style == "title":
        return 1
    return 0


def _clean_heading(text):
    """Normalisasi judul bab: '4.0 Results:' -> 'results'."""
    t = text.replace("\u00a0", " ").replace("\t", " ").strip()
    # penomoran bertingkat: 4 / 4. / 4.0 / 4.1.2 / IV. / (4)
    t = re.sub(r"^\s*(?:chapter|bab)?\s*\(?(?:\d+|[IVXLivxl]+)(?:[\.\-]\d+)*[\.\)]?\s+", "", t)
    t = re.sub(r"^\s*(?:\d+|[IVXLivxl]+)[\.\)]\s*", "", t)  # sisa '4.' menempel
    t = re.sub(r"[:.;,\-\u2013\u2014]+\s*$", "", t)          # buang ':' '.' di akhir
    t = re.sub(r"\s{2,}", " ", t)
    return t.strip().lower()


def _numbering_prefix(text):
    """
    Ambil prefix penomoran heading, misal '4.1.2 Foo' -> '4.1.2', 'Foo' -> None.
    """
    t = text.replace("\u00a0", " ").strip()
    m = re.match(r"^\(?(\d+(?:\.\d+){0,3})\)?[\.\)]?\s+\S", t)
    if m:
        return m.group(1).rstrip(".")
    return None


# ---------------------------------------------------------------
# Ekstraksi bagian-bagian dokumen
# ---------------------------------------------------------------
def extract_parts(paragraphs, in_table_flags=None):
    """
    ``paragraphs`` harus berupa hasil ``get_all_paragraphs(doc)`` (bukan
    ``doc.paragraphs``) agar konten di dalam tabel ikut terbaca.

    ``in_table_flags`` adalah list boolean sejajar dengan ``paragraphs``
    (elemen kedua dari hasil ``get_all_paragraphs``), dipakai untuk mencegah
    isi sel tabel data (mis. header kolom) dianggap sebagai heading.
    """
    if in_table_flags is None:
        in_table_flags = [False] * len(paragraphs)
    parts = {
        "title": None, "abstract": None, "keywords_raw": None,
        "headings": [], "references": [],
        # indeks paragraf (untuk anchor comment) — merujuk posisi di list `paragraphs`
        "title_idx": None, "abstract_idx": None, "keywords_idx": None,
        "ref_heading_idx": None, "ref_idxs": [], "heading_idxs": [],
        "heading_styles": [],   # level heading versi gaya Word (0/1/2/3), sejajar dgn heading_idxs
        "heading_raw": [],      # teks asli heading (belum dinormalisasi), sejajar dgn heading_idxs
        "heading_numpr": [],    # True jika heading pakai automatic list numbering Word, sejajar dgn heading_idxs
        "short_lines": [],
        "doc_word_start_idx": None, "doc_word_end_idx": None,
    }

    texts = [(i, p.text.strip(), p) for i, p in enumerate(paragraphs) if p.text.strip()]

    # Label-label bagian depan yang BUKAN judul (sering muncul di kolom tabel
    # terpisah, mis. kolom kiri berisi logo + "Article History:" yang secara
    # posisi/urutan XML bisa lebih dulu daripada kolom kanan berisi judul).
    FRONT_MATTER_LABELS = [
        r"^article\s+history", r"^corresponding\s+author", r"^received\b",
        r"^revised\b", r"^accepted\b", r"^available\s+online", r"^how\s+to\s+cite",
        r"^p-?issn", r"^e-?issn", r"^\d{4}[\-–]\d{4}$", r"^doi\s*[:.]?", r"^©",
        r"^vol\.?\s*\d+", r"^no\.?\s*\d+",
    ]

    # --- Title: style 'Title', atau paragraf non-kosong pertama yang layak jadi judul ---
    for i, t, p in texts:
        style = (p.style.name or "").lower() if p.style else ""
        if style == "title":
            parts["title"] = t
            parts["title_idx"] = i
            break
    if parts["title"] is None:
        for i, t, p in texts:
            tl = t.lower().strip()
            if _wc(t) < 3:
                continue
            if _match_any(tl, FRONT_MATTER_LABELS):
                continue
            parts["title"] = t
            parts["title_idx"] = i
            break
    if parts["title"] is None and texts:
        parts["title"] = texts[0][1]
        parts["title_idx"] = texts[0][0]

    # --- Abstract & Keywords ---
    abs_start = None
    for idx, (i, t, p) in enumerate(texts):
        if re.match(r"^abstra(?:ct|k)\b", t.lower()):
            abs_start = idx
            break
    if abs_start is not None:
        buf = []
        first = texts[abs_start][1]
        inline = re.sub(r"^abstra(?:ct|k)\s*[:.\-—]?\s*", "", first, flags=re.I).strip()
        parts["abstract_idx"] = texts[abs_start][0]
        if inline:
            buf.append(inline)
        for idx in range(abs_start + 1, len(texts)):
            t = texts[idx][1]
            if re.match(r"^(keywords?|kata\s+kunci)\b", t.lower()):
                break
            if _is_headingish(texts[idx][2], in_table_flags[texts[idx][0]]) and \
               _match_any(_clean_heading(t), [r"introduction", r"pendahuluan"]):
                break
            buf.append(t)
        parts["abstract"] = " ".join(buf).strip() or None

    for i, t, p in texts:
        m = re.match(r"^(keywords?|kata\s+kunci)\s*[:.\-—]?\s*(.*)$", t, flags=re.I)
        if m:
            parts["keywords_raw"] = m.group(2).strip()
            parts["keywords_idx"] = i
            break

    # --- Headings + kandidat baris pendek (fallback) ---
    for i, t, p in texts:
        if _is_headingish(p, in_table_flags[i]):
            parts["headings"].append(_clean_heading(t))
            parts["heading_idxs"].append(i)
            parts["heading_styles"].append(_heading_style_level(p))
            parts["heading_raw"].append(t)
            parts["heading_numpr"].append(_has_auto_numbering(p))
        elif _wc(t) <= 8:
            parts["short_lines"].append(_clean_heading(t))
        # Tambahan: kadang penulis menempelkan heading bab khusus
        # (Acknowledgement/Author Contribution) di TENGAH paragraf lain lewat
        # manual line break (Shift+Enter), bukan sebagai paragraf baru yang
        # sesungguhnya — sehingga baris heading itu tidak pernah terlihat oleh
        # pengecekan di atas (karena keseluruhan paragraf terlalu panjang).
        # Di sini paragraf dipecah per baris internal (\n) dan tiap baris
        # PENDEK diperiksa juga sebagai kandidat heading, supaya kasus
        # semacam ini tetap terdeteksi.
        if "\n" in t:
            for sub in t.split("\n"):
                sub_clean = sub.strip(" \t")
                if sub_clean and sub_clean != t and _wc(sub_clean) <= 8:
                    parts["short_lines"].append(_clean_heading(sub_clean))

    # --- References ---
    ref_idx = None
    for pos, (i, t, p) in enumerate(texts):
        if _is_headingish(p, in_table_flags[i]) and _match_any(t.lower(), REFERENCE_HEADINGS):
            ref_idx = pos
        elif re.match(r"^(references?|daftar\s+pustaka)\s*$", t, flags=re.I):
            ref_idx = pos
    if ref_idx is not None:
        parts["ref_heading_idx"] = texts[ref_idx][0]
    if ref_idx is not None:
        stop_after = [r"appendix", r"lampiran", r"acknowledg", r"author", r"biograph"]
        for pos in range(ref_idx + 1, len(texts)):
            t = texts[pos][1]
            p = texts[pos][2]
            if _is_headingish(p, in_table_flags[texts[pos][0]]) and _match_any(_clean_heading(t), stop_after):
                break
            if _wc(t) >= 4:  # buang nomor halaman / artefak pendek
                parts["references"].append(t)
                parts["ref_idxs"].append(texts[pos][0])

    # --- Rentang untuk hitung total kata artikel (Abstract s.d. akhir References) ---
    if parts["abstract_idx"] is not None:
        parts["doc_word_start_idx"] = parts["abstract_idx"]
    elif texts:
        parts["doc_word_start_idx"] = texts[0][0]
    if parts["ref_idxs"]:
        parts["doc_word_end_idx"] = parts["ref_idxs"][-1]
    elif texts:
        parts["doc_word_end_idx"] = texts[-1][0]

    return parts


# ---------------------------------------------------------------
# Klasifikasi referensi (heuristik)
# ---------------------------------------------------------------
def _extract_journal_name(ref):
    """
    Ambil segmen nama jurnal dari string referensi APA standar, mis.
    'Journal of Public Health and Pharmacy' dari:
    '...Focused on Patient Satisfaction: The Case of Indonesia. Journal of
    Public Health and Pharmacy, 5(2), 297-307. doi:...'

    Dipakai supaya kata bahasa/negara Indonesia yang muncul di JUDUL artikel
    (mis. "The Case of Indonesia", "...di Indonesia") tidak keliru dianggap
    sebagai penanda bahwa JURNALNYA sendiri berbahasa/terbitan Indonesia.
    Nama jurnal ada tepat sebelum pola 'Vol(Issue), halaman' pada referensi
    artikel jurnal.
    """
    m = _find_journal_segment(ref)
    if m:
        return m.group(1).strip()
    return None


def classify_reference(ref):
    """
    Mengklasifikasikan jenis referensi berdasarkan POLA STRUKTURAL metadata
    bibliografi secara menyeluruh (penulis, tahun, judul, nama jurnal,
    volume/issue, halaman, DOI) — BUKAN semata-mata dari kemunculan kata
    tertentu seperti "Journal" atau "Proceedings". Dengan begitu, referensi
    yang formatnya sudah menyerupai artikel jurnal standar APA (mis. "...
    Advances in Neural Information Processing Systems, 33, 1877-1901.")
    tetap terklasifikasi & tervalidasi dengan benar meskipun tidak
    mengandung kata "journal" secara harfiah.
    """
    t = ref.lower()

    # Sinyal STRUKTURAL utama: apakah referensi memiliki pola bibliografi
    # khas artikel berkala ("Nama Berkala, Vol(Issue), Halaman" atau
    # "Nama Berkala, Vol, Halaman") — ini jauh lebih andal daripada mencari
    # kata "journal" secara harfiah, karena mencerminkan struktur sitasi APA
    # yang sebenarnya (nama berkala + penomoran + halaman).
    has_journal_structure = _find_journal_segment(ref) is not None
    # Sinyal kata kunci pendukung (mis. ada kata "journal"/"quarterly", atau
    # pola "12(3)"/"pp. 45-67"/"vol./no." yang longgar) — tetap dipakai
    # sebagai sinyal SEKUNDER untuk referensi yang polanya kurang standar
    # (mis. tanpa titik pemisah yang rapi) tapi tetap punya penanda jurnal.
    has_journal_keyword = _match_any(t, JOURNAL_CUES)

    is_proc = _match_any(t, PROC_MARKERS)
    is_preprint = _match_any(t, PREPRINT_MARKERS)
    is_web = _match_any(t, WEB_MARKERS) or (("http" in t or "www." in t) and not re.search(r"doi", t))
    is_book = _match_any(t, BOOK_MARKERS)

    # Penanda "referensi berbahasa/terbitan Indonesia" HANYA dicek dari nama
    # jurnalnya (bila bisa diisolasi) — BUKAN dari judul artikel. Tanpa ini,
    # judul artikel berbahasa Inggris yang menyebut "Indonesia" (mis. studi
    # kasus di Indonesia) bisa keliru menurunkan jurnal internasional
    # menjadi "Jurnal Nasional".
    journal_name = _extract_journal_name(ref)
    if journal_name:
        has_id = _match_any(journal_name.lower(), ID_MARKERS)
    else:
        has_id = _match_any(t, ID_MARKERS)

    has_journal_cue = has_journal_structure or has_journal_keyword

    # Prosiding/konferensi tetap diprioritaskan bila ada kata kuncinya secara
    # eksplisit (mis. "Proceedings of ..."), meskipun formatnya juga
    # menyerupai jurnal (banyak prosiding dikutip dengan format Vol,Halaman).
    if is_proc:
        return "Prosiding/Konferensi"
    if has_journal_cue and not has_id:
        return "Jurnal Internasional"
    if has_journal_cue and has_id:
        return "Jurnal Nasional"
    # Preprint (arXiv/SSRN/dll.) dicek SETELAH pola jurnal terstruktur, supaya
    # artikel yang SUDAH terbit di jurnal (dan kebetulan juga menyebut versi
    # arXiv-nya di teks referensi) tetap diklasifikasikan sebagai jurnal.
    if is_preprint:
        return "Preprint (arXiv/SSRN/dll.)"
    if is_book:
        return "Buku"
    if is_web:
        return "Website/Laporan"
    if has_id:
        return "Sumber Indonesia (non-jurnal)"
    return "Lainnya/Tidak terdeteksi"


def _is_span_italic(paragraph, ref_text, start, end):
    """
    True/False jika rentang karakter [start:end) pada ``ref_text`` (mis. nama
    jurnal) sebagian besar (>=70%) diformat italic pada paragraf Word aslinya.
    Mengembalikan None jika tidak dapat diverifikasi (paragraf tidak
    tersedia/tidak cocok) — kondisi None TIDAK dianggap sebagai kesalahan,
    supaya tidak muncul false positive saat pemetaan teks gagal.
    """
    if paragraph is None or not getattr(paragraph, "runs", None):
        return None
    raw = paragraph.text
    offset = raw.find(ref_text)
    if offset == -1:
        snippet = ref_text[:20]
        offset = raw.find(snippet) if snippet else -1
        if offset == -1:
            return None
    abs_start, abs_end = offset + start, offset + end
    pos = 0
    covered = italic_covered = 0
    for run in paragraph.runs:
        r_start, r_end = pos, pos + len(run.text)
        pos = r_end
        ov_s, ov_e = max(abs_start, r_start), min(abs_end, r_end)
        if ov_s < ov_e:
            covered += (ov_e - ov_s)
            if run.italic:
                italic_covered += (ov_e - ov_s)
    if covered == 0:
        return None
    return (italic_covered / covered) >= 0.7


# ---------------------------------------------------------------
# Validasi rinci per-referensi terhadap format APA 6th Edition
# ---------------------------------------------------------------
# Pola langsung "Volume(Issue), Halaman" atau "Volume, Halaman" — dipakai
# oleh _pages_present() TANPA bergantung pada berhasilnya ekstraksi nama
# jurnal, supaya referensi yang nama jurnalnya gagal terdeteksi (atau yang
# tidak memakai nomor issue sama sekali, mis. jurnal continuous-pagination
# seperti Decision Support Systems) tetap dapat diperiksa halamannya.
PAGES_AFTER_VOL_PATTERNS = [
    # Vol(Issue), rentang halaman — mis. "12(3), 45-67"
    re.compile(r"\d{1,4}\s*\(\d+[a-zA-Z]?\)\s*,\s*(?:pp\.?\s*)?\d{1,5}\s*[\-\u2013]\s*\d{1,5}\b"),
    # Vol(Issue), halaman/nomor artikel tunggal — mis. "11(18), 2498"
    re.compile(r"\d{1,4}\s*\(\d+[a-zA-Z]?\)\s*,\s*(?:pp\.?\s*)?[a-zA-Z]?\d{1,8}\b"),
    # Vol saja (tanpa issue), rentang halaman — mis. "12, 45-67"
    re.compile(r"\b\d{1,4}\s*,\s*(?:pp\.?\s*)?\d{1,5}\s*[\-\u2013]\s*\d{1,5}\b"),
    # Vol saja (tanpa issue), halaman/nomor artikel tunggal — mis. "176, 114051"
    re.compile(r"\b\d{1,4}\s*,\s*(?:pp\.?\s*)?[a-zA-Z]?\d{1,8}\b"),
]


def _pages_present(ref):
    """
    True jika referensi mencantumkan halaman artikel — baik berupa RENTANG
    (mis. "297-307") ATAU nomor halaman/artikel TUNGGAL (mis. "2498" atau
    "114051", umum pada jurnal seperti Healthcare/MDPI, BMC, Decision
    Support Systems, dsb.) — dengan mencari pola "Vol(Issue), Halaman" atau
    "Vol, Halaman" langsung pada teks. Deteksi ini SENGAJA dibuat tidak
    bergantung pada berhasilnya ekstraksi nama jurnal (_find_journal_segment),
    supaya referensi yang nama jurnalnya sulit diisolasi tetap bisa diperiksa
    halamannya secara independen.

    Pencarian dibatasi pada bagian SEBELUM doi/URL, supaya digit di dalam DOI
    (mis. "10.1186/s12913-025-13172-z" yang mengandung "025-13172") tidak
    keliru terbaca sebagai rentang/nomor halaman.
    """
    cut = re.search(r"\bdoi\b|https?://|10\.\d{4,9}/", ref, re.I)
    scope = ref[:cut.start()] if cut else ref
    return any(p.search(scope) for p in PAGES_AFTER_VOL_PATTERNS)



def validate_reference_format(ref, ref_type, paragraph=None):
    """
    Memeriksa SATU referensi secara rinci terhadap format APA 6th Edition dan
    mengembalikan SEMUA kekurangan yang ditemukan (bukan hanya kesalahan
    pertama) sebagai list of dict {"code": ..., "message": ...}. Setiap
    ``message`` sudah memuat penjelasan spesifik + rekomendasi perbaikan,
    supaya pengguna tidak perlu mencari sendiri letak kesalahannya.

    ``paragraph`` (opsional) adalah objek paragraf python-docx tempat
    referensi ini berada — dipakai untuk memeriksa formatting (mis. italic
    pada nama jurnal) yang tidak bisa dideteksi dari teks polos saja.

    Validasi ini bersifat estimasi otomatis berbasis pola teks & formatting
    Word; hasil akhir tetap perlu diverifikasi oleh editor.
    """
    issues = []
    ref = (ref or "").strip()
    if not ref:
        return issues
    is_article = ref_type in ("Jurnal Internasional", "Jurnal Nasional")

    # 1) Tahun publikasi dalam format APA "(YYYY)" — atau, untuk sumber
    # web/berita tanpa tahun terbit yang jelas, penanda "(n.d.)" yang juga
    # sah menurut APA 6th Edition dan karena itu TIDAK di-flag sebagai
    # kekurangan.
    year_m = YEAR_RE.search(ref)
    if not year_m and not NO_DATE_RE.search(ref):
        issues.append({
            "code": "year_missing",
            "message": ("Tahun publikasi tidak ditemukan. APA 6th Edition mensyaratkan tahun "
                        "publikasi ditulis dalam tanda kurung tepat setelah nama penulis, contoh: "
                        "Smith, J. A. (2021). Mohon tambahkan tahun publikasi pada referensi ini."),
            "message_en": ("Publication year not found. APA 6th Edition requires the publication "
                           "year to appear in parentheses immediately after the author name, for "
                           "example: Smith, J. A. (2021). Please add the publication year to this "
                           "reference."),
        })

    # 2) Format nama penulis
    # PENTING: hanya buang whitespace di ujung, JANGAN buang '.' di ujung —
    # titik tersebut biasanya bagian dari inisial penulis (mis. "Budi, S."),
    # sehingga jika dibuang akan salah menganggap format inisial keliru.
    # Untuk referensi "(n.d.)" (lihat NO_DATE_RE), posisi penanda tanggal
    # dipakai sebagai pengganti year_m supaya bagian nama penulis tetap bisa
    # diperiksa alih-alih langsung dianggap "tidak dapat diverifikasi".
    nd_m = None if year_m else NO_DATE_RE.search(ref)
    date_marker_start = year_m.start() if year_m else (nd_m.start() if nd_m else None)
    author_part = ref[:date_marker_start].strip() if date_marker_start is not None else None
    if author_part is not None:
        if author_part and not AUTHOR_START_RE.match(author_part) \
           and not ORG_AUTHOR_RE.match(author_part) \
           and not SINGLE_NAME_AUTHOR_RE.match(author_part):
            preview = author_part[:80] + ("…" if len(author_part) > 80 else "")
            issues.append({
                "code": "author_format",
                "message": ('Nama penulis ("{}") tampaknya belum sesuai format APA 6th Edition '
                            "(seharusnya: NamaBelakang, Inisial. — misal 'Smith, J. A.'). Mohon "
                            "periksa urutan nama, tanda koma, dan penggunaan inisial pada seluruh "
                            "penulis referensi ini.").format(preview),
                "message_en": ('The author name ("{}") does not appear to follow APA 6th Edition '
                               "format (it should be: Surname, Initials. — for example "
                               "'Smith, J. A.'). Please check the name order, commas, and use of "
                               "initials for every author in this reference.").format(preview),
            })
    else:
        issues.append({
            "code": "author_unverified",
            "message": ("Format nama penulis belum dapat diverifikasi otomatis karena tahun "
                        "publikasi tidak terdeteksi pada referensi ini. Mohon periksa manual apakah "
                        "nama penulis sudah mengikuti format APA (NamaBelakang, Inisial.)."),
            "message_en": ("The author name format could not be verified automatically because no "
                           "publication year was detected in this reference. Please check manually "
                           "whether the author name follows APA format (Surname, Initials.)."),
        })

    # 3) Judul artikel & nama jurnal (segmen setelah tahun/penanda tanggal)
    title_text = journal_name = journal_span = None
    date_marker_end = year_m.end() if year_m else (nd_m.end() if nd_m else None)
    if date_marker_end is not None:
        rest = ref[date_marker_end:]
        jm = _find_journal_segment(rest)
        if jm:
            title_text = rest[:jm.start()].strip(" .")
            journal_name = jm.group(1).strip()
            journal_span = (date_marker_end + jm.start(1), date_marker_end + jm.end(1))
        else:
            stripped = rest.strip(" .")
            # Untuk referensi non-jurnal (web/laporan), potong bagian judul
            # SEBELUM frasa baku APA 6th Edition seperti "Retrieved ... from
            # ..." atau URL/DOI — bagian tersebut BUKAN bagian dari judul,
            # jadi tidak boleh ikut dinilai saat memeriksa Title Case vs
            # sentence case (sebelumnya seluruh sisa teks, termasuk klausa
            # retrieval ini, salah ikut dianggap judul).
            cut_m = re.search(r"\.\s*(?:Retrieved\b|In\s+[A-Z]|http|www\.)", stripped)
            if cut_m:
                stripped = stripped[:cut_m.start()]
            elif re.match(r"^\s*Retrieved\b", stripped, re.I):
                # Referensi tanpa judul tersendiri, mis. "Author. (n.d.).
                # Retrieved from URL" — tidak ada judul untuk diperiksa.
                stripped = ""
            title_text = stripped.strip(" .") or None

    if title_text:
        words = title_text.split()
        significant = [w for w in words if len(re.sub(r"[^A-Za-z]", "", w)) > 3]
        if len(significant) >= 3:
            cap_ratio = sum(1 for w in significant if w[:1].isupper()) / len(significant)
            if cap_ratio > 0.6:
                preview = title_text[:80] + ("…" if len(title_text) > 80 else "")
                issues.append({
                    "code": "title_case",
                    "message": ('Judul artikel ("{}") tampak ditulis dengan huruf kapital di setiap '
                                "kata (Title Case), padahal APA 6th Edition mensyaratkan sentence "
                                "case untuk judul artikel (hanya huruf pertama kalimat/sub-judul dan "
                                "nama diri yang kapital). Mohon sesuaikan kapitalisasi judul.").format(preview),
                    "message_en": ('The article title ("{}") appears to be written in Title Case, '
                                   "whereas APA 6th Edition requires sentence case for article "
                                   "titles (only the first letter of the title and subtitle, plus "
                                   "proper nouns, are capitalised). Please adjust the "
                                   "capitalisation of the title.").format(preview),
                })
    elif year_m and is_article:
        issues.append({
            "code": "title_missing",
            "message": ("Judul artikel tidak terdeteksi setelah tahun publikasi. Mohon pastikan "
                        "referensi mengikuti urutan format APA: Penulis. (Tahun). Judul artikel. "
                        "Nama Jurnal, Volume(Issue), Halaman."),
            "message_en": ("No article title was detected after the publication year. Please make "
                           "sure the reference follows the APA order: Author. (Year). Article "
                           "title. Journal Name, Volume(Issue), Pages."),
        })

    # 4)–6) Nama jurnal (italic), volume/issue, dan halaman — khusus artikel jurnal
    if is_article:
        if not journal_name:
            issues.append({
                "code": "journal_name_missing",
                "message": ("Nama jurnal tidak dapat diidentifikasi dari teks referensi ini, "
                            "meskipun referensi terklasifikasi sebagai artikel jurnal. Mohon "
                            "pastikan nama jurnal dicantumkan tepat sebelum volume(issue) dan "
                            "halaman, sesuai format APA."),
                "message_en": ("The journal name could not be identified from this reference, even "
                               "though the reference was classified as a journal article. Please "
                               "make sure the journal name appears immediately before the "
                               "volume(issue) and page numbers, following APA format."),
            })
        else:
            italic = _is_span_italic(paragraph, ref, *journal_span) if journal_span else None
            if italic is False:
                issues.append({
                    "code": "journal_not_italic",
                    "message": ('Nama jurnal ("{}") belum ditulis dengan format cetak miring '
                                "(italic). APA 6th Edition mewajibkan nama jurnal beserta nomor "
                                "volume ditulis italic. Mohon ubah format teks tersebut menjadi "
                                "italic.").format(journal_name),
                    "message_en": ('The journal name ("{}") is not italicised. APA 6th Edition '
                                   "requires the journal name and the volume number to be set in "
                                   "italics. Please change the formatting of that text to "
                                   "italic.").format(journal_name),
                })

        vol_m = VOL_ISSUE_RE.search(ref)
        if not vol_m:
            issues.append({
                "code": "volume_issue_missing",
                "message": ("Nomor volume dan/atau issue belum dicantumkan. Format APA untuk "
                            "artikel jurnal mensyaratkan penomoran seperti '12(3)' setelah nama "
                            "jurnal. Mohon tambahkan nomor volume(issue) pada referensi ini."),
                "message_en": ("The volume and/or issue number is missing. APA format for journal "
                               "articles requires numbering such as '12(3)' after the journal "
                               "name. Please add the volume(issue) number to this reference."),
            })

        if not _pages_present(ref):
            issues.append({
                "code": "pages_missing",
                "message": ("Nomor halaman artikel tidak ditemukan. Mohon tambahkan halaman "
                            "artikel — boleh berupa rentang halaman (contoh: 297-307) atau "
                            "nomor halaman/artikel tunggal (contoh: 2498) — sesuai format APA."),
                "message_en": ("No page numbers were found for the article. Please add the article "
                               "pages, either as a page range (for example 297-307) or as a single "
                               "page/article number (for example 2498), following APA format."),
            })

    # 7) DOI — kewajiban, keberadaan URL sebagai pengganti, dan validitas format
    if needs_doi(ref_type):
        doi_m = DOI_PATTERN.search(ref)
        if not doi_m:
            if URL_NOT_DOI_RE.search(ref):
                issues.append({
                    "code": "doi_missing_url_used",
                    "message": ("DOI belum dicantumkan — referensi ini menggunakan tautan URL biasa "
                                "sebagai gantinya. Artikel jurnal umumnya memiliki DOI resmi (format "
                                "https://doi.org/10.xxxx/xxxxx). Mohon periksa apakah artikel ini "
                                "memiliki DOI resmi (misalnya melalui crossref.org atau situs jurnal "
                                "terkait) dan gunakan format tersebut, bukan tautan biasa."),
                    "message_en": ("No DOI is provided; this reference uses a plain URL instead. "
                                   "Journal articles normally have an official DOI in the format "
                                   "https://doi.org/10.xxxx/xxxxx. Please check whether this "
                                   "article has an official DOI (for example via crossref.org or "
                                   "the journal website) and use that format rather than a plain "
                                   "link."),
                })
            else:
                issues.append({
                    "code": "doi_missing",
                    "message": ("DOI tidak dicantumkan. Referensi ini terdeteksi sebagai artikel "
                                "jurnal namun belum mencantumkan DOI. Mohon tambahkan DOI dengan "
                                "format https://doi.org/10.xxxx/xxxxx."),
                    "message_en": ("No DOI is provided. This reference was detected as a journal "
                                   "article but does not include a DOI. Please add the DOI in the "
                                   "format https://doi.org/10.xxxx/xxxxx."),
                })
        else:
            doi_str = doi_m.group(1)
            trimmed = doi_str.rstrip(".,;)") if doi_m.end(1) == len(ref) else doi_str
            suffix = trimmed.split("/", 1)[1] if "/" in trimmed else ""
            invalid_reason = invalid_reason_en = None
            if len(suffix) < 3:
                invalid_reason = "tampak tidak lengkap/terpotong"
                invalid_reason_en = "appears incomplete or truncated"
            elif len(trimmed) > 60:
                invalid_reason = ("tampak tercampur dengan teks lain (tidak ada spasi pemisah "
                                  "setelah DOI) atau terlalu panjang untuk pola DOI standar")
                invalid_reason_en = ("appears to run together with other text (no separating space "
                                     "after the DOI) or is too long for a standard DOI pattern")
            elif not re.match(r"^10\.\d{4,9}/\S+$", trimmed):
                invalid_reason = "tidak sesuai pola standar 10.xxxx/xxxxx"
                invalid_reason_en = "does not match the standard 10.xxxx/xxxxx pattern"
            if invalid_reason:
                issues.append({
                    "code": "doi_invalid",
                    "message": ('DOI yang tercantum ("{}") {}. Mohon periksa kembali dan pastikan '
                                "DOI ditulis lengkap, benar, dan terpisah dari teks lain.").format(
                                    trimmed, invalid_reason),
                    "message_en": ('The DOI provided ("{}") {}. Please re-check it and make sure '
                                   "the DOI is complete, correct, and separated from surrounding "
                                   "text.").format(trimmed, invalid_reason_en),
                })

    return issues, journal_name


# ---------------------------------------------------------------
# Ekstraksi sitasi in-text (gaya APA) — dipakai untuk validasi sitasi pada
# bagian Methodology, dan untuk validasi dua-arah sitasi <-> daftar pustaka.
# ---------------------------------------------------------------
# Token nama penulis: nama biasa (diawali huruf kapital), TERMASUK nama
# dengan prefiks huruf kecil + apostrof yang lazim pada nama Prancis/Belanda
# (mis. "d'Astous", "l'Hritier", "o'Brien" ditulis "O'Brien" juga tercakup
# karena huruf kapital plain sudah cocok pada cabang kedua). Karakter
# lanjutan memakai kelas Unicode "huruf apa pun" (bukan hanya Latin-1),
# supaya nama dengan huruf non-Latin-1 (mis. "Kıymalıoğlu" berbahasa Turki)
# tidak terpotong jadi hanya huruf awalnya saja.
AUTHOR_TOKEN = r"(?:[a-z]{1,3}['\u2019])?[A-Z](?:[^\W\d_]|['\u2019\-])*"

# Sitasi naratif: "Smith (2020) menyatakan...", "Smith & Jones (2020)...",
# "Smith et al. (2020)...", "d'Astous (2021)..."
NARRATIVE_CITE_RE = re.compile(
    r"\b(" + AUTHOR_TOKEN + r")"
    r"(?:\s+(?:&|and|dan|et\s+al\.?)\s+" + AUTHOR_TOKEN + r")?"
    r"\s*\(((?:19|20)\d{2})[a-z]?\)"
)

# Prefiks non-nama yang kadang mendahului sitasi dalam kurung dan perlu
# dilewati sebelum mencari token nama penulis, mis. "(lihat Smith, 2020)".
_CITE_LEAD_STRIP_RE = re.compile(
    r"^\s*(?:e\.g\.,?\s*|see\s+|cf\.\s*|dalam\s+|in\s+|lihat\s+)", re.I)


def _extract_citation_keys_with_pos(paragraphs, end_idx):
    """
    Menelusuri ``paragraphs[0:end_idx]`` (badan artikel, TIDAK termasuk
    daftar pustaka) dan mengekstrak seluruh sitasi in-text gaya APA — baik
    dalam kurung, mis. "(Smith, 2020)", "(Smith & Jones, 2020; Lee, 2019)",
    maupun naratif, mis. "Smith (2020) menyatakan...".

    Untuk gaya dalam-kurung, sitasi HARUS mengikuti bentuk baku APA
    "(Penulis, Tahun)" — yaitu ADA TANDA KOMA tepat sebelum tahun. Syarat ini
    sengaja diberlakukan supaya frasa lain yang kebetulan memuat angka
    4-digit di dalam kurung TIDAK ikut terbaca sebagai sitasi, mis. rentang
    tanggal "(February-April 2025)" atau nomor hibah "(Award No.: .../2025,
    Date: 28 July 2025)" — keduanya tidak punya koma tepat sebelum angka
    tahun, sehingga tidak akan cocok.

    Nama penulis diambil sebagai token PERTAMA pada awal isi kurung (bukan
    hasil pencarian bebas di tengah teks), supaya nama dengan prefiks huruf
    kecil + apostrof (mis. "d'Astous") tidak terpotong menjadi "Astous".

    Mengembalikan dict {(surname_lower, year_digits): first_paragraph_idx}
    — tahun disimpan HANYA 4 digit (mengabaikan sufiks huruf disambiguasi
    "a"/"b", karena sufiks ini kerap dipakai tidak konsisten antara badan
    teks & daftar pustaka) supaya pencocokan dengan referensi lebih toleran.
    """
    positions = {}
    for i, p in enumerate(paragraphs):
        if i >= end_idx:
            break
        text = p.text
        if not text.strip():
            continue
        for pm in re.finditer(r"\(([^()]{3,300})\)", text):
            group = pm.group(1)
            for part in re.split(r";\s*", group):
                ym = re.search(r",\s*((?:19|20)\d{2})[a-z]?\b", part)
                if not ym:
                    continue
                before = _CITE_LEAD_STRIP_RE.sub("", part[:ym.start()]).strip()
                am = re.match(AUTHOR_TOKEN, before)
                if am:
                    key = (am.group(0).lower(), ym.group(1))
                    positions.setdefault(key, i)
        for nm in NARRATIVE_CITE_RE.finditer(text):
            key = (nm.group(1).lower(), nm.group(2))
            positions.setdefault(key, i)
    return positions


def _extract_ref_citation_key(ref):
    """Ambil (surname_lower, year_digits) dari penulis PERTAMA & tahun suatu
    entri referensi, untuk dicocokkan dengan sitasi in-text. Mengembalikan
    None jika tahun atau nama penulis tidak dapat diidentifikasi (referensi
    semacam ini tidak diikutsertakan dalam pengecekan dua-arah supaya tidak
    menimbulkan false positive).

    Jika referensi tidak memiliki tahun terbit eksplisit tapi ditandai
    "(n.d.)" (lazim pada sumber web/berita tanpa tanggal terbit yang jelas,
    mis. halaman Google Maps atau media sosial), year_digits dikembalikan
    sebagai literal string "nd" — pemanggil (validate_citation_reference_
    match) lalu mencocokkan referensi semacam ini dengan sitasi in-text
    berdasarkan NAMA PENULIS SAJA, mengabaikan tahun, karena penulis kerap
    menyitasi sumber "n.d." memakai tahun akses/retrieval yang tidak selalu
    sama dengan tahun pada entri daftar pustaka."""
    year_m = YEAR_RE.search(ref)
    if year_m:
        author_part = ref[:year_m.start()].strip()
        am = re.match(AUTHOR_TOKEN, author_part)
        if not am:
            return None
        year_digits_m = re.match(r"(19|20)\d{2}", year_m.group(1))
        if not year_digits_m:
            return None
        return am.group(0).lower(), year_digits_m.group(0)

    nd_m = NO_DATE_RE.search(ref)
    if nd_m:
        author_part = ref[:nd_m.start()].strip()
        am = re.match(AUTHOR_TOKEN, author_part)
        if am:
            return am.group(0).lower(), "nd"
    return None


# ---------------------------------------------------------------
# Validasi sitasi pendukung pada bagian Methodology/Research Methodology
# ---------------------------------------------------------------
def _chapter_paragraph_range(parts, chapter_label, total_len):
    """
    Mengembalikan (start_idx, end_idx, heading_idx) berupa rentang indeks
    paragraf (dari list hasil ``get_all_paragraphs``) yang menjadi ISI bab
    ``chapter_label`` (sesuai definisi CHAPTERS) — mencakup seluruh
    subsection di bawahnya, sampai tepat sebelum heading bab utama
    BERIKUTNYA (atau sebelum References/akhir dokumen jika bab ini yang
    terakhir). Mengembalikan None jika heading bab tersebut tidak ditemukan.
    """
    headings = parts["headings"]
    heading_idxs = parts["heading_idxs"]
    n = len(headings)
    if n == 0:
        return None

    chapter_positions = {}
    for sec_no, (label, pats) in enumerate(CHAPTERS, start=1):
        for pos, h in enumerate(headings):
            if _match_any(h, pats):
                chapter_positions[pos] = sec_no
                break

    target_pos = None
    for pos, sec_no in chapter_positions.items():
        if CHAPTERS[sec_no - 1][0] == chapter_label:
            target_pos = pos
            break
    if target_pos is None:
        return None

    sorted_positions = sorted(chapter_positions.keys())
    order = sorted_positions.index(target_pos)
    end_pos = sorted_positions[order + 1] if order + 1 < len(sorted_positions) else None

    start_idx = heading_idxs[target_pos] + 1
    end_idx = heading_idxs[end_pos] if end_pos is not None else (parts.get("ref_heading_idx") or total_len)
    return start_idx, end_idx, heading_idxs[target_pos]


def validate_methodology_citations(parts, paragraphs):
    """
    Memeriksa apakah bagian Methodology/Research Methodology memiliki
    minimal satu sitasi pendukung (mis. untuk metode, teknik analisis,
    instrumen, atau pendekatan penelitian yang dipakai). Mengembalikan dict
    {"idx", "heading", "n_citations"} atau None jika bagian Methodology
    tidak ditemukan sama sekali (sudah ditangani oleh cek "Struktur 5 Bab
    Utama" secara terpisah).
    """
    rng = _chapter_paragraph_range(parts, "Research Methodology", len(paragraphs))
    if rng is None:
        return None
    start_idx, end_idx, heading_idx = rng
    end_idx = min(end_idx, len(paragraphs))
    keys = _extract_citation_keys_with_pos(paragraphs[start_idx:end_idx], end_idx - start_idx)
    return {
        "idx": heading_idx,
        "heading": parts["heading_raw"][parts["heading_idxs"].index(heading_idx)].strip(),
        "n_citations": len(keys),
    }


# ---------------------------------------------------------------
# Validasi dua-arah: sitasi in-text <-> entri daftar pustaka
# ---------------------------------------------------------------
def validate_citation_reference_match(parts, paragraphs, ref_rows):
    """
    Memastikan setiap referensi pada daftar pustaka disitasi minimal satu
    kali di badan teks, DAN setiap sitasi di badan teks memiliki entri yang
    sesuai pada daftar pustaka.

    Mengembalikan (uncited_refs, orphan_citations):
      - uncited_refs: list baris ref_rows yang tidak pernah disitasi
        (hanya utk referensi yang penulis+tahunnya berhasil diparse —
        referensi yang tidak bisa diparse TIDAK di-flag, untuk menghindari
        false positive).
      - orphan_citations: list dict {"surname", "year", "idx"} — sitasi
        in-text yang tidak memiliki pasangan pada daftar pustaka.

    Referensi "n.d." (lihat _extract_ref_citation_key) dicocokkan dengan
    sitasi in-text berdasarkan NAMA PENULIS SAJA (mengabaikan tahun),
    karena tahun yang dipakai di badan teks untuk sumber semacam ini
    (mis. tahun akses situs web) tidak selalu identik dengan penanda
    "(n.d.)" pada daftar pustaka.

    Pemeriksaan ini bersifat estimasi otomatis berbasis pola teks (sitasi
    APA dalam kurung maupun naratif) — hasil akhir tetap perlu diverifikasi
    oleh editor.
    """
    end = parts.get("ref_heading_idx")
    if end is None:
        end = len(paragraphs)
    positions = _extract_citation_keys_with_pos(paragraphs, end)
    in_text_keys = set(positions.keys())
    in_text_surnames = {surname for surname, _ in in_text_keys}

    uncited_refs = []
    ref_keys_all = set()
    nd_surnames = set()
    for row in ref_rows:
        key = _extract_ref_citation_key(row["text"])
        if key is None:
            continue
        surname, year = key
        if year == "nd":
            nd_surnames.add(surname)
            cited = surname in in_text_surnames
        else:
            ref_keys_all.add(key)
            cited = key in in_text_keys
        if not cited:
            uncited_refs.append(row)

    orphan_citations = []
    for key in sorted(in_text_keys - ref_keys_all):
        surname, year = key
        if surname in nd_surnames:
            continue
        orphan_citations.append({"surname": surname, "year": year, "idx": positions[key]})

    return uncited_refs, orphan_citations


# ---------------------------------------------------------------
# Validasi jumlah kata total artikel (Abstract s.d. References)
# ---------------------------------------------------------------
def count_article_words(paragraphs, parts):
    start = parts.get("doc_word_start_idx")
    end = parts.get("doc_word_end_idx")
    if start is None or end is None:
        return 0
    total = 0
    for i, p in enumerate(paragraphs):
        if start <= i <= end:
            total += _wc(p.text)
    return total


# ---------------------------------------------------------------
# Validasi struktur hierarki heading (Section / Subsection / Sub-subsection)
# ---------------------------------------------------------------
def validate_heading_hierarchy(parts):
    """
    Memeriksa apakah heading di bawah tiap bab utama (Section) sudah mengikuti
    hierarki Section/Subsection/Sub-subsection yang konsisten sesuai template
    jurnal — TANPA mengasumsikan atau menyarankan format penomoran tertentu
    (mis. '4.1', '4.1.1'), karena penomoran berbeda-beda antar template.

    Validasi hanya berlaku untuk heading yang benar-benar merupakan bagian
    dari struktur dokumen (heading bergaya Word "Heading"/"Title", atau
    baris pendek full-bold di luar tabel) — bukan caption Table/Figure/
    Formula maupun isi sel tabel data, yang sudah disaring lebih dulu di
    ``extract_parts``/``_is_headingish``.

    Mengembalikan list of dict:
      {"idx": <indeks paragraf untuk anchor comment>,
       "current": <teks heading apa adanya>,
       "message": <pesan lengkap untuk comment, tanpa saran penomoran>}
    """
    issues = []
    headings = parts["headings"]
    heading_idxs = parts["heading_idxs"]
    heading_raw = parts["heading_raw"]
    heading_numpr = parts.get("heading_numpr") or [False] * len(headings)
    n = len(headings)
    if n == 0:
        return issues

    # Tentukan posisi tiap bab utama (Section) berdasarkan CHAPTERS,
    # nomor section mengikuti urutan baku template (1..5).
    chapter_positions = {}  # posisi index-dalam-headings -> nomor section
    for sec_no, (label, pats) in enumerate(CHAPTERS, start=1):
        for pos, h in enumerate(headings):
            if _match_any(h, pats):
                chapter_positions[pos] = sec_no
                break

    if not chapter_positions:
        return issues

    sorted_positions = sorted(chapter_positions.keys())

    for order, pos in enumerate(sorted_positions):
        sec_no = chapter_positions[pos]
        start_pos = pos + 1
        end_pos = sorted_positions[order + 1] if order + 1 < len(sorted_positions) else n

        for hp in range(start_pos, end_pos):
            text_raw = heading_raw[hp]
            text_clean = headings[hp]
            prefix = _numbering_prefix(text_raw)

            # Lewati heading yang sebenarnya adalah bab khusus/References
            # yang terdeteksi ikut di rentang (mis. Acknowledgement duluan
            # sebelum References) — jangan divalidasi sebagai subsection.
            if _match_any(text_clean, [p for _, pats in SPECIAL_CHAPTERS for p in pats]) or \
               _match_any(text_clean, REFERENCE_HEADINGS):
                continue

            if prefix:
                # Sudah ada penomoran eksplisit (tertulis di teks) -> anggap
                # sudah mengikuti hierarki template, tidak perlu di-flag.
                continue

            if heading_numpr[hp]:
                # Heading ini memakai automatic list numbering Word (numPr) —
                # nomor seperti "2.5" dirender otomatis oleh Word walau tidak
                # tersimpan sebagai teks pada paragraf. Anggap sudah bernomor
                # dengan benar, tidak perlu di-flag.
                continue

            # Heading di bawah bab utama ini tidak punya penomoran berjenjang
            # sama sekali -> struktur Section/Subsection/Sub-subsection-nya
            # belum jelas/konsisten dengan template jurnal. Sistem TIDAK
            # menyarankan nomor tertentu (mis. '4.1' atau '4.1.2') karena
            # format penomoran berbeda-beda antar template jurnal.
            issues.append({
                "idx": heading_idxs[hp],
                "current": text_raw.strip(),
                "message": (
                    "STRUKTUR HEADING: Heading \u201c{}\u201d belum mengikuti struktur "
                    "Section/Subsection/Sub-subsection yang sesuai dengan template "
                    "jurnal. Silakan sesuaikan penomoran dan hierarki heading ini "
                    "dengan format Section, Subsection, dan Sub-subsection pada "
                    "template jurnal yang digunakan.".format(text_raw.strip())
                ),
                "message_en": (
                    "HEADING STRUCTURE: The heading \u201c{}\u201d does not yet follow the "
                    "Section/Subsection/Sub-subsection structure used in the journal "
                    "template. Please adjust the numbering and hierarchy of this heading "
                    "to match the Section, Subsection, and Sub-subsection format of the "
                    "journal template in use.".format(text_raw.strip())
                ),
            })

    return issues


# ---------------------------------------------------------------
# Validasi subsection wajib Bab 5 (Conclusion) sesuai template Goodwood:
# 5.1 Conclusion, 5.2 Research Limitations,
# 5.3 Suggestions and Directions for Future Research
# ---------------------------------------------------------------
def validate_conclusion_subsections(parts):
    """
    Template jurnal Goodwood mewajibkan Bab 5 dipecah menjadi tiga subsection:
    5.1 Conclusion, 5.2 Research Limitations, dan 5.3 Suggestions and
    Directions for Future Research. Fungsi ini memeriksa apakah subsection
    "Research Limitations" dan "Suggestions and Directions for Future
    Research" sudah ada sebagai heading tersendiri di bawah Bab 5.

    Subsection "Conclusion" dianggap otomatis terpenuhi oleh heading Bab 5
    itu sendiri (mis. "5. Conclusion" atau "5.1 Conclusion"), sehingga tidak
    perlu divalidasi keberadaannya secara terpisah.

    Mengembalikan list of dict (kosong jika Bab 5 tidak ditemukan sama
    sekali, atau jika kedua subsection wajib sudah ada):
      {"idx": <indeks paragraf heading Bab 5, untuk anchor comment>,
       "current": <teks heading Bab 5 apa adanya>,
       "missing": <list nama subsection yang belum ditemukan>,
       "message": <pesan lengkap untuk comment>}
    """
    issues = []
    headings = parts["headings"]
    heading_idxs = parts["heading_idxs"]
    heading_raw = parts["heading_raw"]
    n = len(headings)
    if n == 0:
        return issues

    # Cari posisi heading Bab 5 (Conclusion) berdasarkan daftar CHAPTERS.
    concl_pos = None
    for label, pats in CHAPTERS:
        if label != "Conclusion":
            continue
        for pos, h in enumerate(headings):
            if _match_any(h, pats):
                concl_pos = pos
                break
        break

    if concl_pos is None:
        # Bab 5 (Conclusion) tidak ditemukan sama sekali — sudah ditangani
        # oleh pengecekan "Struktur 5 Bab Utama", jadi tidak perlu di-flag lagi di sini.
        return issues

    # Subsection "Conclusion" otomatis terpenuhi oleh heading Bab 5 itu sendiri.
    required = CONCLUSION_SUBSECTIONS[1:]  # Research Limitations, Suggestions & Future Research
    found = set()

    for hp in range(concl_pos + 1, n):
        text_clean = headings[hp]
        # Berhenti begitu masuk ke bab khusus (Acknowledgement/Author
        # Contribution) atau References — di luar cakupan Bab 5.
        if _match_any(text_clean, REFERENCE_HEADINGS) or \
           _match_any(text_clean, [p for _, pats in SPECIAL_CHAPTERS for p in pats]):
            break
        for sub_label, sub_pats in required:
            if sub_label not in found and _match_any(text_clean, sub_pats):
                found.add(sub_label)

    missing = [label for label, _ in required if label not in found]
    if not missing:
        return issues

    required_list = ", ".join(
        "5.{} {}".format(i, label) for i, (label, _) in enumerate(CONCLUSION_SUBSECTIONS, start=1)
    )
    missing_list = ", ".join(missing)

    issues.append({
        "idx": heading_idxs[concl_pos],
        "current": heading_raw[concl_pos].strip(),
        "missing": missing,
        "message": (
            "STRUKTUR BAB 5 (CONCLUSION): Berdasarkan template jurnal, Bab 5 "
            "harus dipecah menjadi subsection {}. Subsection yang belum "
            "terdeteksi pada manuskrip ini: {}. Mohon strukturkan ulang Bab 5 "
            "sesuai format tersebut.".format(required_list, missing_list)
        ),
        "message_en": (
            "CHAPTER 5 STRUCTURE (CONCLUSION): According to the journal template, "
            "Chapter 5 must be divided into the subsections {}. The following "
            "subsections were not detected in this manuscript: {}. Please "
            "restructure Chapter 5 accordingly.".format(required_list, missing_list)
        ),
    })

    return issues
# ---------------------------------------------------------------
# Validasi cross-reference Table & Figure — memastikan setiap tabel/gambar
# disebut/dirujuk di badan teks pembahasan, bukan hanya muncul sebagai
# caption tanpa kalimat pengantar.
# ---------------------------------------------------------------
def validate_cross_references(paragraphs):
    """
    Memeriksa setiap caption "Table N" / "Figure N" (beserta padanan
    Indonesia "Tabel N" / "Gambar N") dan mengecek apakah elemen tersebut
    juga disebut di tempat LAIN pada badan teks (mis. "As shown in Table 2,",
    "Figure 3 illustrates...", "Dari Tabel 2, terlihat bahwa..."). Jika
    penyebutan elemen tersebut HANYA muncul pada captionnya sendiri (tidak
    pernah dirujuk di kalimat pembahasan), elemen ini dianggap belum
    dirujuk dan akan di-flag.

    ``paragraphs`` adalah hasil ``get_all_paragraphs(doc)`` (list paragraf,
    termasuk isi tabel), supaya caption yang kebetulan diletakkan sebagai
    baris judul di dalam tabel juga ikut terdeteksi.

    Mengembalikan list of dict:
      {"idx": <indeks paragraf caption, untuk anchor comment>,
       "type": "Table" | "Figure",
       "number": <nomor elemen, string>,
       "caption": <cuplikan teks caption>,
       "message": <pesan comment>}
    """
    issues = []
    all_text = " \n ".join(p.text for p in paragraphs if p.text.strip())
    seen = set()

    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        for elem_type, words, pats in CROSSREF_ELEMENTS:
            matched_num = None
            for pat in pats:
                m = re.match(pat, text, re.I)
                if m:
                    matched_num = m.group(1)
                    break
            if matched_num is None:
                continue
            key = (elem_type, matched_num)
            if key in seen:
                break
            seen.add(key)

            # Hitung kemunculan "Table N" / "Tabel N" (atau "Figure N" /
            # "Fig. N" / "Gambar N") di SELURUH dokumen, termasuk caption
            # itu sendiri. Jika totalnya hanya 1 (yaitu caption itu sendiri),
            # berarti elemen ini tidak pernah dirujuk di badan teks.
            word_alt = "|".join(re.escape(w) for w in words)
            mention_pat = re.compile(
                r"\b(?:" + word_alt + r")\.?\s*" + re.escape(matched_num) + r"\b", re.I)
            count = len(mention_pat.findall(all_text))

            if count <= 1:
                elem_lower = "tabel" if elem_type == "Table" else "gambar"
                issues.append({
                    "idx": i,
                    "type": elem_type,
                    "number": matched_num,
                    "caption": text[:120],
                    "message": (
                        "CROSS-REFERENCE: {et} {num} belum dirujuk di badan teks pembahasan "
                        "artikel — kemunculannya hanya ditemukan pada caption ini. Setiap {el} "
                        "sebaiknya diperkenalkan dengan kalimat pengantar di badan teks sebelum "
                        "ditampilkan, misalnya \u201cAs shown in {et} {num}, ...\u201d, \u201cFrom "
                        "{et} {num}, ...\u201d, atau \u201c{et} {num} illustrates...\u201d. Mohon "
                        "tambahkan kalimat rujukan yang sesuai pada bagian pembahasan "
                        "terkait.".format(et=elem_type, num=matched_num, el=elem_lower)
                    ),
                    "message_en": (
                        "CROSS-REFERENCE: {et} {num} is not referred to anywhere in the body "
                        "text of the article; the only occurrence found is this caption itself. "
                        "Every {ell} should be introduced by a lead-in sentence in the body text "
                        "before it is displayed, for example \u201cAs shown in {et} {num}, "
                        "...\u201d, \u201cFrom {et} {num}, ...\u201d, or \u201c{et} {num} "
                        "illustrates...\u201d. Please add an appropriate referring sentence in "
                        "the related discussion section.".format(
                            et=elem_type, num=matched_num, ell=elem_type.lower())
                    ),
                })
            break

    return issues


def screen_docx(path):
    doc = Document(path)
    result, _, _ = screen_document(doc)
    return result


def screen_document(doc):
    all_paragraphs, in_table_flags = get_all_paragraphs(doc)
    parts = extract_parts(all_paragraphs, in_table_flags)
    checks = []

    # 1) Title
    title = parts["title"] or ""
    n_title = _wc(title)
    checks.append({
        "name": "Judul (maks. {} kata)".format(TITLE_MAX_WORDS),
        "passed": bool(title) and n_title <= TITLE_MAX_WORDS,
        "detail": ('"{}" — {} kata'.format(title, n_title)) if title else "Judul tidak terdeteksi.",
        "title": title,
        "n_words": n_title,
    })

    # 2) Keywords
    kw_raw = parts["keywords_raw"]
    if kw_raw is not None:
        kws = [k.strip() for k in re.split(r"[;,•·]", kw_raw) if k.strip()]
        ok = KEYWORDS_MIN <= len(kws) <= KEYWORDS_MAX
        detail = "{} keywords terdeteksi: {}".format(len(kws), "; ".join(kws))
    else:
        kws, ok = [], False
        detail = "Bagian Keywords tidak ditemukan."
    checks.append({
        "name": "Keywords ({}–{})".format(KEYWORDS_MIN, KEYWORDS_MAX),
        "passed": ok, "detail": detail,
        "keywords": kws if kw_raw is not None else None,
    })

    # 2b) Format Keywords — Capital Each Word (Title Case)
    # Hanya relevan jika keyword benar-benar ditemukan; jika tidak, kekurangan
    # itu sudah dilaporkan oleh cek "Keywords" di atas, sehingga tidak perlu
    # duplikasi comment.
    if kws:
        bad_kws = [k for k in kws if not _is_title_case_keyword(k)]
        ok_case = len(bad_kws) == 0
        kw_case_detail = (
            "Seluruh keyword sudah menggunakan format Capital Each Word (Title Case)."
            if ok_case else
            "Keyword berikut belum menggunakan format Capital Each Word (Title Case): {}.".format(
                "; ".join('"{}"'.format(k) for k in bad_kws))
        )
        checks.append({
            "name": "Format Keywords (Title Case)",
            "passed": ok_case, "detail": kw_case_detail,
            "bad_keywords": bad_kws,
        })

    # 3) Abstract
    abstract = parts["abstract"]
    if abstract:
        n_abs = _wc(abstract)
        ok = ABSTRACT_MIN <= n_abs <= ABSTRACT_MAX
        detail = "{} kata (syarat: {}–{} kata).".format(n_abs, ABSTRACT_MIN, ABSTRACT_MAX)
        if n_abs < ABSTRACT_MIN:
            detail += " Kurang {} kata.".format(ABSTRACT_MIN - n_abs)
        elif n_abs > ABSTRACT_MAX:
            detail += " Lebih {} kata.".format(n_abs - ABSTRACT_MAX)
    else:
        ok, detail = False, "Bagian Abstract tidak ditemukan."
    checks.append({
        "name": "Abstrak ({}–{} kata)".format(ABSTRACT_MIN, ABSTRACT_MAX),
        "passed": ok, "detail": detail,
        "word_count": _wc(abstract) if abstract else 0,
    })

    # 4) 5 bab utama
    headings = parts["headings"]
    chapter_rows = []
    all_ch = True
    short_lines = parts["short_lines"]
    for label, pats in CHAPTERS:
        found = next((h for h in headings if _match_any(h, pats)), None)
        if not found:
            found = next((h for h in short_lines if _match_any(h, pats)), None)
        chapter_rows.append({"label": label, "found": found})
        if not found:
            all_ch = False
    checks.append({
        "name": "Struktur 5 Bab Utama",
        "passed": all_ch,
        "detail": "chapters",  # dirender khusus di template
        "chapters": chapter_rows,
    })

    # 5) Special chapters
    sp_rows, all_sp = [], True
    for label, pats in SPECIAL_CHAPTERS:
        found = next((h for h in headings if _match_any(h, pats)), None)
        if not found:
            found = next((h for h in short_lines if _match_any(h, pats)), None)
        sp_rows.append({"label": label, "found": found})
        if not found:
            all_sp = False
    checks.append({
        "name": "Bab Khusus (Acknowledgement & Author Contribution)",
        "passed": all_sp, "detail": "chapters", "chapters": sp_rows,
    })

    # 6) Referensi — jumlah, komposisi internasional, dan DOI
    refs = parts["references"]
    ref_idxs = parts["ref_idxs"]
    ref_rows = []
    for i, r in enumerate(refs):
        rtype = classify_reference(r)
        para = all_paragraphs[ref_idxs[i]] if i < len(ref_idxs) else None
        issues, journal_name = validate_reference_format(r, rtype, para)
        ref_rows.append({
            "no": i + 1, "text": r, "type": rtype,
            "has_doi": has_doi(r),
            "needs_doi": needs_doi(rtype),
            "issues": issues,
            "apa_ok": len(issues) == 0,
            "journal_name": journal_name,
        })
    n_ref = len(refs)
    n_intl = sum(1 for r in ref_rows if r["type"] == "Jurnal Internasional")
    pct = (100.0 * n_intl / n_ref) if n_ref else 0.0
    ok_count = n_ref >= REF_MIN
    ok_pct = pct >= REF_INTL_PCT
    checks.append({
        "name": "Jumlah Referensi (min. {})".format(REF_MIN),
        "passed": ok_count,
        "detail": "{} referensi terdeteksi.".format(n_ref) if n_ref else
                  "Bagian References tidak ditemukan atau kosong.",
        "n_refs": n_ref,
    })
    checks.append({
        "name": "Jurnal Internasional (min. {:.0f}%)".format(REF_INTL_PCT),
        "passed": ok_pct and n_ref > 0,
        "detail": "{} dari {} referensi ({:.1f}%) terdeteksi sebagai jurnal internasional. "
                  "Klasifikasi ini bersifat estimasi otomatis — mohon verifikasi pada tabel di bawah."
                  .format(n_intl, n_ref, pct) if n_ref else "Tidak ada referensi untuk dianalisis.",
        "references": ref_rows,
        "pct": round(pct, 1),
        "n_refs": n_ref,
        "n_intl": n_intl,
    })

    # 6b) Validasi DOI referensi
    doi_expected = [r for r in ref_rows if r["needs_doi"]]
    doi_missing = [r for r in doi_expected if not r["has_doi"]]
    if n_ref == 0:
        doi_detail = "Tidak ada referensi untuk diperiksa."
        doi_ok = False
    elif not doi_expected:
        doi_detail = ("Tidak ada referensi yang teridentifikasi sebagai artikel jurnal, "
                      "sehingga pemeriksaan DOI tidak berlaku.")
        doi_ok = True
    elif not doi_missing:
        doi_detail = "Seluruh {} referensi jurnal telah mencantumkan DOI.".format(len(doi_expected))
        doi_ok = True
    else:
        doi_detail = ("{} dari {} referensi jurnal belum mencantumkan DOI. "
                      "Nomor: {}.").format(
                          len(doi_missing), len(doi_expected),
                          ", ".join(str(r["no"]) for r in doi_missing))
        doi_ok = False
    checks.append({
        "name": "Validasi DOI Referensi",
        "passed": doi_ok,
        "detail": doi_detail,
        "doi_missing": doi_missing,
    })

    # 6c) Validasi rinci format APA 6th Edition per-referensi (satu per satu,
    # menampilkan SELURUH kekurangan yang ditemukan, bukan hanya yang pertama)
    n_apa_issue_refs = sum(1 for r in ref_rows if not r["apa_ok"])
    if n_ref == 0:
        apa_detail = "Tidak ada referensi untuk diperiksa."
        apa_ok = False
    elif n_apa_issue_refs == 0:
        apa_detail = "Seluruh {} referensi sudah sesuai format APA 6th Edition.".format(n_ref)
        apa_ok = True
    else:
        apa_detail = ("{} dari {} referensi memiliki kekurangan format APA 6th Edition. "
                      "Lihat rincian kekurangan & rekomendasi perbaikan per referensi di bawah.").format(
                          n_apa_issue_refs, n_ref)
        apa_ok = False
    checks.append({
        "name": "Validasi Format APA per Referensi",
        "passed": apa_ok,
        "detail": apa_detail,
        "ref_apa_rows": ref_rows,
    })

    # 6d) Validasi sitasi pendukung pada bagian Methodology/Research Methodology
    method_result = validate_methodology_citations(parts, all_paragraphs)
    if method_result is None:
        method_detail = ("Bagian Methodology/Research Methodology tidak ditemukan pada dokumen "
                          "(lihat kriteria \u201cStruktur 5 Bab Utama\u201d).")
        method_passed = False
    elif method_result["n_citations"] > 0:
        method_detail = ("Bagian Methodology memiliki {} sitasi pendukung terdeteksi (estimasi "
                          "otomatis berdasarkan pola sitasi APA).").format(method_result["n_citations"])
        method_passed = True
    else:
        method_detail = ("Bagian Methodology/Research Methodology belum memiliki satu pun sitasi "
                          "pendukung. Umumnya bagian ini perlu menyitasi sumber yang mendukung "
                          "metode, teknik analisis, instrumen, atau pendekatan penelitian yang "
                          "digunakan. Mohon tambahkan sitasi yang relevan.")
        method_passed = False
    checks.append({
        "name": "Sitasi pada Bagian Methodology",
        "passed": method_passed,
        "detail": method_detail,
        "methodology_result": method_result,
    })

    # 6e) Validasi dua-arah sitasi in-text <-> entri daftar pustaka
    uncited_refs, orphan_citations = validate_citation_reference_match(parts, all_paragraphs, ref_rows)
    n_uncited, n_orphan = len(uncited_refs), len(orphan_citations)
    if n_ref == 0:
        citmatch_detail = "Tidak ada referensi untuk diperiksa kesesuaiannya dengan sitasi."
        citmatch_passed = False
    elif n_uncited == 0 and n_orphan == 0:
        citmatch_detail = ("Seluruh referensi yang berhasil diverifikasi telah disitasi di badan "
                          "teks, dan seluruh sitasi terdeteksi memiliki pasangan pada daftar pustaka.")
        citmatch_passed = True
    else:
        msg_parts = []
        if n_uncited:
            msg_parts.append("{} referensi tidak pernah disitasi di badan teks (No. {})".format(
                n_uncited, ", ".join(str(r["no"]) for r in uncited_refs)))
        if n_orphan:
            sample = ", ".join("{} ({})".format(o["surname"].capitalize(), o["year"])
                                for o in orphan_citations[:15])
            if n_orphan > 15:
                sample += ", dst."
            msg_parts.append("{} sitasi di badan teks tidak memiliki pasangan pada daftar pustaka: {}".format(
                n_orphan, sample))
        citmatch_detail = ("; ".join(msg_parts) + ". Pemeriksaan bersifat estimasi otomatis "
                          "berdasarkan pola sitasi APA (dalam kurung maupun naratif) — mohon "
                          "verifikasi manual.")
        citmatch_passed = False
    checks.append({
        "name": "Kesesuaian Sitasi & Daftar Pustaka",
        "passed": citmatch_passed,
        "detail": citmatch_detail,
        "uncited_refs": uncited_refs,
        "orphan_citations": orphan_citations,
    })

    # 6f) Validasi minimal penggunaan referensi dari jurnal Goodwood Publishing
    goodwood_matches = validate_goodwood_journal_usage(ref_rows)
    n_goodwood = len(goodwood_matches)
    goodwood_ok = n_goodwood >= GOODWOOD_JOURNAL_MIN
    if n_ref == 0:
        goodwood_detail = "Tidak ada referensi untuk diperiksa."
    elif goodwood_ok:
        goodwood_detail = ("{} referensi dari jurnal terbitan Goodwood Publishing terdeteksi "
                          "(syarat minimal {}).").format(n_goodwood, GOODWOOD_JOURNAL_MIN)
    else:
        goodwood_detail = ("Artikel belum memenuhi persyaratan minimum penggunaan referensi dari "
                          "jurnal Goodwood Publishing: baru {} dari minimal {} referensi jurnal "
                          "Goodwood yang terdeteksi. Mohon tambahkan referensi dari jurnal-jurnal "
                          "terbitan Goodwood Publishing.").format(n_goodwood, GOODWOOD_JOURNAL_MIN)
    checks.append({
        "name": "Referensi Jurnal Goodwood (min. {})".format(GOODWOOD_JOURNAL_MIN),
        "passed": goodwood_ok and n_ref > 0,
        "detail": goodwood_detail,
        "goodwood_matches": goodwood_matches,
        "goodwood_count": n_goodwood,
        "goodwood_min": GOODWOOD_JOURNAL_MIN,
        "n_refs": n_ref,
    })

    # 7) Jumlah kata total artikel (Abstract s.d. References)
    total_words = count_article_words(all_paragraphs, parts)
    ok_words = DOC_WORDS_MIN <= total_words <= DOC_WORDS_MAX
    if total_words == 0:
        w_detail = "Jumlah kata tidak dapat dihitung (Abstract/References tidak terdeteksi)."
    else:
        w_detail = "{} kata (dihitung dari Abstract s.d. References; syarat: {:,}\u2013{:,} kata).".format(
            total_words, DOC_WORDS_MIN, DOC_WORDS_MAX).replace(",", ".")
        if total_words < DOC_WORDS_MIN:
            w_detail += " Kurang {} kata.".format(DOC_WORDS_MIN - total_words)
        elif total_words > DOC_WORDS_MAX:
            w_detail += " Lebih {} kata.".format(total_words - DOC_WORDS_MAX)
    checks.append({
        "name": "Jumlah Kata Artikel ({:,}\u2013{:,} kata)".format(DOC_WORDS_MIN, DOC_WORDS_MAX).replace(",", "."),
        "passed": ok_words,
        "detail": w_detail,
        "word_count": total_words,
    })

    # 8) Struktur hierarki heading (Section/Subsection/Sub-subsection)
    hierarchy_issues = validate_heading_hierarchy(parts)
    checks.append({
        "name": "Struktur Penomoran Section/Subsection",
        "passed": len(hierarchy_issues) == 0,
        "detail": ("Seluruh subsection/sub-subsection sudah diberi penomoran berjenjang sesuai template."
                   if not hierarchy_issues else
                   "{} heading belum mengikuti format penomoran berjenjang (lihat rincian & comment di dokumen).".format(
                       len(hierarchy_issues))),
        "hierarchy_issues": hierarchy_issues,
    })

    # 9) Subsection wajib Bab 5 (5.1 Conclusion, 5.2 Research Limitations,
    #    5.3 Suggestions and Directions for Future Research)
    conclusion_issues = validate_conclusion_subsections(parts)
    checks.append({
        "name": "Struktur Subsection Bab 5 (Conclusion)",
        "passed": len(conclusion_issues) == 0,
        "detail": ("Bab 5 sudah terstruktur menjadi 5.1 Conclusion, 5.2 Research Limitations, "
                   "dan 5.3 Suggestions and Directions for Future Research."
                   if not conclusion_issues else
                   "Bab 5 belum dipecah menjadi subsection wajib: {} (lihat comment di dokumen).".format(
                       ", ".join(conclusion_issues[0]["missing"]))),
        "conclusion_issues": conclusion_issues,
    })

    # 10) Cross-reference Table & Figure — setiap tabel/gambar harus dirujuk
    #     dengan kalimat pengantar di badan teks pembahasan, bukan hanya
    #     muncul sebagai caption.
    crossref_issues = validate_cross_references(all_paragraphs)
    n_tables = sum(1 for i in crossref_issues if i["type"] == "Table")
    n_figs = sum(1 for i in crossref_issues if i["type"] == "Figure")
    checks.append({
        "name": "Cross-reference Table & Figure",
        "passed": len(crossref_issues) == 0,
        "detail": ("Seluruh Table dan Figure sudah dirujuk dengan kalimat pengantar di badan teks."
                   if not crossref_issues else
                   "{} Table dan {} Figure belum dirujuk di badan teks pembahasan (lihat rincian & "
                   "comment di dokumen).".format(n_tables, n_figs)),
        "crossref_issues": crossref_issues,
    })

    passed = sum(1 for c in checks if c["passed"])
    result = {
        "checks": checks,
        "passed": passed,
        "total": len(checks),
        "verdict": "LOLOS SCREENING" if passed == len(checks) else "PERLU REVISI",
        "title": title,
    }
    return result, parts, all_paragraphs
